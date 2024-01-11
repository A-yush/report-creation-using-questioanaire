# importing libraries

import json
import boto3
import time
import os,binascii
from datetime import *
from docx import Document
from docx.shared import Inches
from io import BytesIO

# initializing dynamoDB and S3 boto3 clinets
ddbClient = boto3.client('dynamodb')
S3Client = boto3.client('s3')

'''

sample input event format

{
    "question1":"Who owns the enterprise data strategy in your organization",
    "question2":"How do you ensure that you have the right data and code between all your environments",
    "answer1":"No clear ownership for the enterprise data strategy today",
    "answer2":"We have strict version control policies for data and code promotion and they are followed most of the time"
}

Input variables for your AWS resources

1. bucketName: Replace with your S3 bucket where documents will be stored
2. TableName: replace table name with your DynamoDB table where questions, answers,recommendations and other metadata is stored.

'''

# function to get the recommendations from dynamoDB and formatting it for document file
def getRecommendation(question,answer,number):
    answerResponse = ddbClient.get_item(
        TableName = 'dynamoDB-table-name', # Change the table name with your dynamoDB table name
        Key = {
            'question': {'S': question},
            'answer':{'S':answer}
        })
    print(answerResponse)
    rec_title = "=========================================================="  
    question = "Question" + str(number)+ ": " + question
    answer = "Answer "+ str(number)+ ": " + answer
    recommendation = "Recommendation: " + answerResponse["Item"]["recommendation"]["S"]
    score = float(answerResponse["Item"]["score"]["S"])
   
    return question,answer,recommendation, score

# function to create a document, saving docx to s3 and creating s3 presigned url
def create_document(event,number,bucketName):
    document = Document()
    document.add_heading('Recommendations Document', 0)
   
    final_score = 0
    for i in range(1,number+1):
        question_no = "question" + str(i)
        answer_no = "answer" + str(i)
        question = event[question_no]
        answer = event[answer_no]
        ques,ans,recommendation,score = getRecommendation(question,answer,i) # getting recommendation
        final_score += score  # getting final score
        
        # creating document
        q= document.add_paragraph('')
        q.add_run(ques).bold = True  # making question bold
        a= document.add_paragraph('')
        a.add_run(ans).italic = True # making answer italic
        r= document.add_paragraph(recommendation)
        
    score_line = "Cumulated score is: " + str(final_score)
    s= document.add_paragraph('')
    s.add_run(score_line).bold = True # making final score 
    
    # creating s3 key with random number and timestamp
    rand=binascii.b2a_hex(os.urandom(2))
    timenow = datetime.now()
    s3key=str(timenow)+str(rand) +".docx"
    
    # creating a document file and uploading to s3
    with BytesIO() as fileobj:
        document.save(fileobj)
        fileobj.seek(0)
        S3Client.upload_fileobj(fileobj, bucketName, s3key)
        
    
    # creating presigned URL for downloading the .docx file
    presignedurl = S3Client.generate_presigned_url('get_object',Params={'Bucket': bucketName,'Key': s3key},ExpiresIn=600)
    print (presignedurl)
    return presignedurl

# main handler function    
def lambda_handler(event,context):
    jsonInput = json.dumps(event)

    presignedurl = create_document(event,2,"s3-bucket-name") # change the S3 bucket name with your bucket name

    # returning presigned URL of S3 object
    return {
        'statusCode': 200,
        'body': presignedurl
    }
